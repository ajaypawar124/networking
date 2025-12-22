import paramiko
import time
from lxml import etree
from docx import Document
from docx.shared import Pt
from getpass import getpass
from docx.oxml.ns import qn
from datetime import datetime

# ============ CONFIGURATION ============
USERNAME = input("Provide username for devices [default: regress]: ") or "regress"
PASSWORD = getpass("Provide password for devices [enter for regress]: ") or "MaRtInI"

COMMANDS = [
    "show version",
    "show system uptime"
]

DEVICE_FILE = "devices.txt"
OUTPUT_DOCX = f"ssh_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
# =======================================


def ssh_connect(ip, username, password, timeout=10):
    """Establish SSH connection and return client."""
    ssh = paramiko.SSHClient()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(ip, username=username, password=password, timeout=timeout)
    return ssh


def run_commands(ssh, commands):
    """Run multiple commands on an existing SSH session."""
    results = {}
    for cmd in commands:
        try:
            stdin, stdout, stderr = ssh.exec_command(cmd)
            time.sleep(1)
            output = stdout.read().decode(errors="ignore").strip()
            error = stderr.read().decode(errors="ignore").strip()
            if error:
                results[cmd] = f"Error: {error}"
            else:
                results[cmd] = output or "<no output>"
        except Exception as e:
            results[cmd] = f"Failed to execute '{cmd}': {e}"
    return results


def get_device_info(ssh, ip):
    """Fetch hostname and model from XML output."""
    try:
        stdin, stdout, stderr = ssh.exec_command("show version | display xml")
        xml_output = stdout.read().decode(errors="ignore")
        root = etree.fromstring(xml_output.encode())

        hostname = root.xpath("//*[local-name()='host-name']/text()")
        model = root.xpath("//*[local-name()='product-model']/text()")

        name = hostname[0].strip() if hostname else ip
        model_name = model[0].strip() if model else "Unknown"

        return name, model_name
    except Exception:
        return ip, "Unknown"


def main():
    # Load device list
    with open(DEVICE_FILE, "r") as f:
        devices = [line.strip() for line in f if line.strip()]

    doc = Document()
    doc.add_heading("Device Command Outputs", level=0)
    styles = doc.styles
    if "CLIOutput" not in styles:
        cli_style = styles.add_style("CLIOutput", 1)  # 1 = Paragraph style
        cli_font = cli_style.font
        cli_font.name = "Consolas"
        cli_font.size = Pt(10)
        cli_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    summary_data = []
    for ip in devices:
        print(f"\n🔗 Connecting to {ip}...")
        try:
            ssh = ssh_connect(ip, USERNAME, PASSWORD)
            hostname, model = get_device_info(ssh, ip)

            # --- DEVICE HEADING (Heading 1 for TOC/indexing) ---
            heading_text = f"Device: {ip} ({hostname})"
            doc.add_heading(heading_text, level=1)

            # Add Hostname & Model info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Hostname: {hostname}\n").bold = True
            info_para.add_run(f"Model: {model}\n").bold = True

            # --- RUN COMMANDS ---
            results = run_commands(ssh, COMMANDS)
            ssh.close()

            for cmd, output in results.items():
                command_text = f"Command: {cmd} on ({hostname})"
                doc.add_heading(command_text, level=2)
                doc.add_paragraph(output, style="CLIOutput")
            # Separator
            doc.add_paragraph("\n" + "=" * 70 + "\n")

        except Exception as e:
            doc.add_paragraph(f"\nFailed to connect to {ip}: {e}\n")
            print(f"❌ Failed to connect to {ip}: {e}")
            summary_data.append((ip, "Unknown", "No"))
    doc.add_page_break()
    doc.add_heading("Summary of Device Connections", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Index"
    hdr_cells[1].text = "IP Address"
    hdr_cells[2].text = "Hostname"
    hdr_cells[3].text = "Data_Fetched"

    # Add data rows
    for idx, (ip, hostname, status) in enumerate(summary_data, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = ip
        row_cells[2].text = hostname
        row_cells[3].text = status
    # --- SAVE OUTPUT ---
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ All outputs saved to '{OUTPUT_DOCX}'")


if __name__ == "__main__":
    main()

